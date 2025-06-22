import streamlit as st
import os
import io
import tempfile
from PIL import Image
import pandas as pd
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter
import docx
import openpyxl
import pypandoc
import time
import tabula
# Page configuration
st.set_page_config(
    page_title="Universal Convertor",
    page_icon="🍕",
    layout="wide",
    initial_sidebar_state="expanded"
)
# Helper functions for conversions
def pdf_to_word(pdf_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_file:
        pdf_file.write(pdf_bytes)
        pdf_file.flush()
        pdf_file_path = pdf_file.name
    docx_path = pdf_file_path.replace('.pdf', '.docx')
    cv = Converter(pdf_file_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    with open(docx_path, 'rb') as f:
        docx_bytes = f.read()
    # Try to remove temp files, ignore errors if file is still locked
    for path in [pdf_file_path, docx_path]:
        for _ in range(3):
            try:
                os.remove(path)
                break
            except PermissionError:
                time.sleep(0.2)
                continue
            except Exception:
                break
    return docx_bytes

def word_to_pdf(docx_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as docx_file:
        docx_file.write(docx_bytes)
        docx_file.flush()
        pdf_path = docx_file.name.replace('.docx', '.pdf')
        pypandoc.convert_file(docx_file.name, 'pdf', outputfile=pdf_path)
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        os.remove(docx_file.name)
        os.remove(pdf_path)
        return pdf_bytes

def word_to_excel(docx_bytes):
    doc = docx.Document(io.BytesIO(docx_bytes))
    tables = doc.tables
    output = io.BytesIO()
    if tables:
        # If there are tables, extract them
        for idx, table in enumerate(tables):
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            df = pd.DataFrame(data)
            with pd.ExcelWriter(output, engine='openpyxl', mode='a' if idx > 0 else 'w') as writer:
                df.to_excel(writer, index=False, sheet_name=f'Table{idx+1}')
        return output.getvalue()
    else:
        # Fallback: extract paragraphs as before
        data = []
        for para in doc.paragraphs:
            data.append([para.text])
        df = pd.DataFrame(data, columns=['Text'])
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        return output.getvalue()

def excel_to_word(xlsx_bytes):
    df = pd.read_excel(io.BytesIO(xlsx_bytes))
    doc = docx.Document()
    for i, row in df.iterrows():
        doc.add_paragraph(str(row.values))
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def image_convert(image_bytes, output_format):
    image = Image.open(io.BytesIO(image_bytes))
    output = io.BytesIO()
    if output_format.upper() == "JPEG" and image.mode == "RGBA":
        image = image.convert("RGB")
    image.save(output, format=output_format)
    return output.getvalue()

def get_pdf_page_count(pdf_bytes):
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        return len(reader.pages)
    except Exception:
        return None

def pdf_to_excel(pdf_bytes, lattice=False, pages='all'):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_file:
        pdf_file.write(pdf_bytes)
        pdf_file.flush()
        pdf_path = pdf_file.name
    try:
        dfs = tabula.read_pdf(
            pdf_path,
            pages=pages,
            multiple_tables=True,
            lattice=lattice,
            stream=not lattice
        )
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for idx, df in enumerate(dfs):
                # Clean DataFrame: convert all values to string and ignore encoding errors
                df = df.applymap(lambda x: str(x).encode('utf-8', errors='ignore').decode('utf-8', errors='ignore') if pd.notnull(x) else '')
                df.to_excel(writer, index=False, sheet_name=f'Table{idx+1}')
        return output.getvalue()
    except Exception as e:
        raise RuntimeError(f'Tabula extraction failed: {e}')
    finally:
        os.remove(pdf_path)

# Streamlit UI
st.title('Universal File Converter')

st.write('Convert between PDF, Word, Excel, and image formats (JPEG, PNG, etc.)')

file = st.file_uploader('Upload a file', type=['pdf', 'docx', 'xlsx', 'jpg', 'jpeg', 'png', 'bmp', 'tiff'])

conversion_types = {
    'PDF to Word': ('pdf', 'docx'),
    'PDF to Excel': ('pdf', 'xlsx'),
    'Word to PDF': ('docx', 'pdf'),
    'Word to Excel': ('docx', 'xlsx'),
    'Excel to Word': ('xlsx', 'docx'),
    'Image to Image': ('image', 'image'),
}

conversion = st.selectbox('Select conversion type', list(conversion_types.keys()))

if file is not None:
    input_ext = file.name.split('.')[-1].lower()
    input_bytes = file.read()
    output_bytes = None
    output_ext = None
    lattice_mode = False
    page_range = 'all'
    page_count = None
    if conversion == 'PDF to Excel' and input_ext == 'pdf':
        lattice_mode = st.checkbox('Use lattice mode for table extraction (better for tables with lines)', value=False)
        page_count = get_pdf_page_count(input_bytes)
        if page_count:
            st.info(f'This PDF has {page_count} page(s).')
        page_range = st.text_input('Page range (e.g. 1, 1-2, 1,3,5 or leave blank for all pages)', value='')
        if not page_range.strip():
            page_range = 'all'
        # Validate page range
        valid_range = True
        if page_range != 'all' and page_count:
            import re
            pages = set()
            for part in page_range.split(','):
                part = part.strip()
                if '-' in part:
                    start, end = part.split('-')
                    try:
                        start, end = int(start), int(end)
                        if start < 1 or end > page_count or start > end:
                            valid_range = False
                        pages.update(range(start, end+1))
                    except:
                        valid_range = False
                else:
                    try:
                        p = int(part)
                        if p < 1 or p > page_count:
                            valid_range = False
                        pages.add(p)
                    except:
                        valid_range = False
            if not valid_range:
                st.error(f'Invalid page range! PDF has {page_count} page(s).')
        if valid_range:
            with st.spinner('Extracting tables from PDF...'):
                try:
                    output_bytes = pdf_to_excel(input_bytes, lattice=lattice_mode, pages=page_range)
                    output_ext = 'xlsx'
                except Exception as e:
                    st.error(f'PDF to Excel conversion failed: {e}')
                    output_bytes = None
    elif conversion == 'PDF to Word' and input_ext == 'pdf':
        output_bytes = pdf_to_word(input_bytes)
        output_ext = 'docx'
    elif conversion == 'Word to PDF' and input_ext == 'docx':
        output_bytes = word_to_pdf(input_bytes)
        output_ext = 'pdf'
    elif conversion == 'Word to Excel' and input_ext == 'docx':
        output_bytes = word_to_excel(input_bytes)
        output_ext = 'xlsx'
    elif conversion == 'Excel to Word' and input_ext == 'xlsx':
        output_bytes = excel_to_word(input_bytes)
        output_ext = 'docx'
    elif conversion == 'Image to Image' and input_ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
        output_format = st.selectbox('Select output image format', ['JPEG', 'PNG', 'BMP', 'TIFF'])
        output_bytes = image_convert(input_bytes, output_format)
        output_ext = output_format.lower()
    else:
        st.error('Invalid file type for selected conversion.')

    if output_bytes:
        st.success('Conversion successful!')
        st.download_button(
            label=f'Download converted file ({output_ext})',
            data=output_bytes,
            file_name=f'converted.{output_ext}'
        ) 
